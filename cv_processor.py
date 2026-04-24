import re
import shutil
import hashlib
import zipfile
import os
import tempfile
import time
from io import BytesIO
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import lxml.etree as etree

# Markers that must exist in the base CV's placeholder paragraphs.
MARKER         = "[ATS_KEYWORDS_HERE]"    # invisible white 1pt injection point
MARKER_VISIBLE = "[ATS_KEYWORDS_VISIBLE]" # visible 3pt white-highlight injection point

_WNS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XMLNS = "http://www.w3.org/XML/1998/namespace"


def _w(local: str) -> str:
    return f"{{{_WNS}}}{local}"


# ── Naming helpers ────────────────────────────────────────────────────────────

def _safe_name(s: str, max_len: int = 40) -> str:
    return re.sub(r'[<>:"/\\|?*\s]+', "_", str(s)).strip("_")[:max_len]


def url_hash(job_url: str) -> str:
    """6-char MD5 prefix of the job URL — makes folder names unique per job."""
    return hashlib.md5(job_url.encode()).hexdigest()[:6]


def job_folder(company: str, title: str, date_str: str, uhash: str = "") -> Path:
    suffix = f"_{uhash}" if uhash else ""
    folder_name = f"{_safe_name(company)}_{_safe_name(title)}_{date_str}{suffix}"
    path = Path.home() / "Desktop" / "job_applications" / folder_name
    path.mkdir(parents=True, exist_ok=True)
    return path


# ── Base CV checks ────────────────────────────────────────────────────────────

def check_base_cv(base_cv_path: str) -> dict:
    """Verify the base CV exists, is readable, and contains both ATS markers."""
    p = Path(base_cv_path)
    result = {
        "exists": p.exists(),
        "file_size": 0,
        "readable": False,
        "has_marker": False,
        "has_marker_visible": False,
        "error": None,
    }
    if not result["exists"]:
        return result
    result["file_size"] = p.stat().st_size
    try:
        doc = Document(str(p))
        result["readable"] = True
        for para in doc.paragraphs:
            if MARKER in para.text:
                result["has_marker"] = True
            if MARKER_VISIBLE in para.text:
                result["has_marker_visible"] = True
    except Exception as e:
        result["error"] = str(e)
    return result


def copy_base_cv_to_temp(base_cv_path: str) -> str:
    """Copy base CV to the local temp directory (away from OneDrive).
    Validates the copy is complete. Returns the temp path."""
    src = Path(base_cv_path)
    if not src.exists():
        raise FileNotFoundError(f"Base CV not found: {base_cv_path}")

    tmp_path = Path(tempfile.gettempdir()) / "gongzuo_base_cv.docx"
    shutil.copy2(str(src), str(tmp_path))

    src_size = src.stat().st_size
    tmp_size = tmp_path.stat().st_size

    if tmp_size == 0:
        raise IOError(
            f"Temp copy is 0 bytes (source was {src_size} bytes). "
            "OneDrive may have the file locked — try again in a moment."
        )
    if tmp_size < src_size * 0.9:
        raise IOError(
            f"Temp copy size mismatch: got {tmp_size} bytes, "
            f"source is {src_size} bytes. Possible partial copy."
        )
    return str(tmp_path)


# ── XML run builders ──────────────────────────────────────────────────────────

def _build_invisible_run(parent, text: str) -> None:
    """Append a white 1pt run (invisible to human readers, readable by ATS)."""
    r = etree.SubElement(parent, _w('r'))
    rPr = etree.SubElement(r, _w('rPr'))

    fonts = etree.SubElement(rPr, _w('rFonts'))
    for attr in ('ascii', 'eastAsia', 'hAnsi', 'cs'):
        fonts.set(_w(attr), 'Times New Roman')

    color = etree.SubElement(rPr, _w('color'))
    color.set(_w('val'), 'FFFFFF')

    sz = etree.SubElement(rPr, _w('sz'));   sz.set(_w('val'), '2')
    szCs = etree.SubElement(rPr, _w('szCs')); szCs.set(_w('val'), '2')

    shd = etree.SubElement(rPr, _w('shd'))
    shd.set(_w('val'), 'clear')
    shd.set(_w('color'), 'auto')
    shd.set(_w('fill'), 'FFFFFF')

    t = etree.SubElement(r, _w('t'))
    t.text = text
    t.set(f"{{{_XMLNS}}}space", 'preserve')


def _build_visible_tiny_run(parent, text: str) -> None:
    """Append a 3pt white-highlight run (tiny but parseable by ATS scanners)."""
    r = etree.SubElement(parent, _w('r'))
    rPr = etree.SubElement(r, _w('rPr'))

    fonts = etree.SubElement(rPr, _w('rFonts'))
    for attr in ('ascii', 'eastAsia', 'hAnsi', 'cs'):
        fonts.set(_w(attr), 'Times New Roman')

    sz = etree.SubElement(rPr, _w('sz'));   sz.set(_w('val'), '6')
    szCs = etree.SubElement(rPr, _w('szCs')); szCs.set(_w('val'), '6')

    hl = etree.SubElement(rPr, _w('highlight'))
    hl.set(_w('val'), 'white')

    t = etree.SubElement(r, _w('t'))
    t.text = text
    t.set(f"{{{_XMLNS}}}space", 'preserve')


# ── Core XML injection ────────────────────────────────────────────────────────

def _inject_in_xml(xml_bytes: bytes, keyword_text: str) -> tuple:
    """
    Parse document.xml, find both ATS marker paragraphs, replace their runs
    with keyword text.  Returns (new_xml_bytes, invisible_found, visible_found).

    Uses lxml for direct XML editing so python-docx never re-saves the document
    (which would alter global formatting properties).
    """
    tree = etree.parse(BytesIO(xml_bytes))
    root = tree.getroot()
    invisible_found = False
    visible_found   = False

    for para in root.iter(_w('p')):
        full = ''.join(t.text or '' for t in para.iter(_w('t')))

        if MARKER in full:
            for r in list(para):
                if r.tag == _w('r'):
                    para.remove(r)
            # Also strip yellow highlight from paragraph-mark formatting
            pPr = para.find(_w('pPr'))
            if pPr is not None:
                pPr_rPr = pPr.find(_w('rPr'))
                if pPr_rPr is not None:
                    hl = pPr_rPr.find(_w('highlight'))
                    if hl is not None:
                        pPr_rPr.remove(hl)
            _build_invisible_run(para, keyword_text)
            invisible_found = True

        elif MARKER_VISIBLE in full:
            for r in list(para):
                if r.tag == _w('r'):
                    para.remove(r)
            _build_visible_tiny_run(para, keyword_text)
            visible_found = True

    out = BytesIO()
    tree.write(out, xml_declaration=True, encoding='UTF-8', standalone=True)
    return out.getvalue(), invisible_found, visible_found


def _rewrite_zip(source_path: str, new_xml: bytes) -> None:
    """Rewrite the docx zip, replacing word/document.xml with new_xml.
    All other zip members are preserved byte-for-byte."""
    tmp = source_path + '.tmp'
    with zipfile.ZipFile(source_path, 'r') as zin:
        with zipfile.ZipFile(tmp, 'w') as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    data = new_xml
                zout.writestr(item, data)
    os.replace(tmp, source_path)


# ── Verification ──────────────────────────────────────────────────────────────

def _verify(output_path: str, keywords: list) -> dict:
    """Reopen the saved file from disk and verify the injection is correct."""
    p = Path(output_path)
    v = {
        "file_exists":    p.exists(),
        "file_size":      p.stat().st_size if p.exists() else 0,
        "marker_gone":    False,
        "keywords_found": False,
        "correct_format": False,
        "passed":         False,
    }
    if not v["file_exists"] or v["file_size"] == 0:
        return v

    doc = Document(output_path)
    all_texts = [para.text for para in doc.paragraphs]

    v["marker_gone"] = not any(
        MARKER in t or MARKER_VISIBLE in t for t in all_texts
    )

    non_empty = [t for t in all_texts if t.strip()]
    if keywords and non_empty:
        v["keywords_found"] = keywords[0].lower() in non_empty[-1].lower()
    else:
        v["keywords_found"] = True

    # Check formatting on the last non-empty paragraph (should be invisible run)
    for para in reversed(doc.paragraphs):
        if para.text.strip():
            runs = para._element.findall(qn('w:r'))
            if runs:
                rPr = runs[0].find(qn('w:rPr'))
                if rPr is not None:
                    sz_el    = rPr.find(qn('w:sz'))
                    color_el = rPr.find(qn('w:color'))
                    sz_val    = sz_el.get(qn('w:val'))    if sz_el    is not None else None
                    color_val = color_el.get(qn('w:val')) if color_el is not None else None
                    v["correct_format"] = (sz_val == '2' and color_val == 'FFFFFF')
            break

    v["passed"] = all([
        v["file_exists"],
        v["file_size"] > 0,
        v["marker_gone"],
        v["keywords_found"],
        v["correct_format"],
    ])
    return v


# ── Main injection ────────────────────────────────────────────────────────────

def inject_ats_keywords(base_cv_path: str, keywords: list, output_path: str) -> dict:
    """
    Copy base CV, inject keywords at both marker paragraphs via direct XML
    editing (zipfile + lxml).  No python-docx save — all original formatting
    is preserved.  Retries once on transient failure.

    Returns a diagnostics dict.  Raises on unrecoverable failure.
    """
    src = Path(base_cv_path)
    if not src.exists():
        raise FileNotFoundError(f"Base CV not found: {base_cv_path}")

    diag = {
        "base_cv_path":     base_cv_path,
        "output_path":      output_path,
        "file_size_source": src.stat().st_size,
        "file_size_output": 0,
        "paragraph_count":  0,
        "marker_found":     False,
        "keywords_count":   len(keywords),
        "keywords":         list(keywords),
        "verification":     {},
        "attempts":         0,
        "error":            None,
    }

    keyword_text = ", ".join(keywords)

    for attempt in range(1, 3):
        diag["attempts"] = attempt
        try:
            shutil.copy2(str(src), output_path)

            with zipfile.ZipFile(output_path, 'r') as z:
                xml_bytes = z.read('word/document.xml')

            new_xml, invisible_found, visible_found = _inject_in_xml(xml_bytes, keyword_text)
            diag["marker_found"] = invisible_found

            if not invisible_found:
                raise ValueError(
                    f"Marker '{MARKER}' not found in the base CV. "
                    "Open the base CV in Word and add the marker text exactly, "
                    f"or re-download the base CV from the repository."
                )

            _rewrite_zip(output_path, new_xml)

            diag["file_size_output"] = Path(output_path).stat().st_size

            # Count paragraphs for diagnostics
            with zipfile.ZipFile(output_path, 'r') as z:
                doc_xml = z.read('word/document.xml')
            diag["paragraph_count"] = doc_xml.count(b'<w:p ')

            v = _verify(output_path, keywords)
            diag["verification"] = v

            if v["passed"]:
                return diag

            failed = [k for k, val in v.items() if k not in ("passed",) and not val]
            if attempt == 2:
                raise RuntimeError(
                    f"CV verification failed after {attempt} attempts. "
                    f"Failed checks: {failed}"
                )
            time.sleep(1)

        except (ValueError, FileNotFoundError):
            raise
        except Exception as e:
            if attempt == 2:
                diag["error"] = str(e)
                raise
            time.sleep(1)

    return diag  # unreachable; satisfies linter


# ── Cover letter document ─────────────────────────────────────────────────────

def save_cover_letter_docx(text: str, output_path: str) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        para = doc.add_paragraph()
        run = para.add_run(block)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(10)

    doc.save(output_path)
    return output_path
