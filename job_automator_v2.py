"""
Job Search Automator v2 — Emilie Quillet
=========================================
What this does every time it runs:
  1. Scrapes LinkedIn, Indeed, Glassdoor for M&A + adjacent roles in Luxembourg
  2. Filters out jobs already seen
  3. For each new job:
      a. Generates a tailored cover letter in Emilie's voice
      b. Injects invisible ATS keywords into a modified copy of her CV (.docx)
      c. Logs everything to a CSV tracker
  4. Saves cover letters as .txt and modified CVs as .docx, organised by date/company

Run manually: python job_automator_v2.py
Schedule:     crontab -e → add: 0 7 * * * python3 /path/to/job_automator_v2.py
"""

import json
import os
import csv
import time
import copy
import re
from datetime import datetime
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jobspy import scrape_jobs


# ─────────────────────────────────────────────────────────────────
# CONFIGURATION — only edit this section
# ─────────────────────────────────────────────────────────────────

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")

# Path to your base CV (.docx) — update this to wherever you saved it
BASE_CV_PATH = Path("CV_Emilie_Quillet_v15.docx")

SEARCH_CONFIG = {
    "search_terms": [
        # Core M&A
        "M&A Analyst",
        "Mergers Acquisitions Analyst",
        "Transaction Advisory",
        "Transaction Services Analyst",
        # Corporate finance / adjacent
        "Corporate Finance Analyst",
        "Financial Analyst Luxembourg",
        "Valuation Analyst",
        "Investment Analyst",
        "FP&A Analyst",
        "Deal Advisory Analyst",
        "Due Diligence Analyst",
        # Consulting adjacent (builds toward M&A)
        "Business Analyst Financial Services",
        "Finance Consultant Luxembourg",
    ],
    "location": "Luxembourg",
    "results_per_search": 15,
    "sites": ["linkedin", "indeed", "glassdoor"],
    "hours_old": 24,
}

# Output folder structure — created automatically on Desktop
OUTPUT_DIR = Path.home() / "Desktop" / "job_applications"
COVER_LETTERS_DIR = OUTPUT_DIR / "cover_letters"
MODIFIED_CVS_DIR = OUTPUT_DIR / "modified_cvs"
SEEN_JOBS_FILE = OUTPUT_DIR / "seen_jobs.json"
CSV_LOG_FILE = OUTPUT_DIR / "job_log.csv"

CSV_HEADERS = [
    "Date Found", "Job Title", "Company", "Location", "Site",
    "Job URL", "Date Posted", "Cover Letter File", "Modified CV File",
    "Applied?", "Response", "Notes"
]


# ─────────────────────────────────────────────────────────────────
# EMILIE'S PROFILE — used by Claude to write cover letters
# ─────────────────────────────────────────────────────────────────

EMILIE_PROFILE = """
CURRENT SITUATION:
- Working as Business Analyst at BearingPoint Luxembourg (April 2026 - present)
  Focus: Banking & Capital Markets consulting, financial transformation, process optimisation,
  client presentations, digital transformation projects in financial sector
- French national, based in Luxembourg, working legally
- Languages: English (native), French (native), Chinese (fluent HSK4)

EDUCATION:
- Master of Financial Engineering, UCLA Anderson School of Management
  GPA: 3.44 | Beta Gamma Sigma Honor Society | Dec 2025
  Key courses: Derivative Markets, Data Analysis & Machine Learning, Managing Finance
- BSc Business Administration Finance, North Carolina State University
  GPA: 4.00 Summa Cum Laude, Dean's List | May 2023
  Key courses: Investments & Portfolio Management, Statistical Programming, Financial Data Analysis
- BBA Global Management Finance, SKEMA Business School
  GPA: 3.87 Summa Cum Laude, Dean's List | Double degree with NCSU | May 2023

EXPERIENCE (most recent first):
1. BearingPoint Luxembourg — Business Analyst (April 2026–present)
   Banking & Capital Markets consulting. Client-facing project work across wealth management,
   private banking, financial transformation. Process optimisation, dashboards, client presentations.

2. Quint Capital Corporation, New York — Financial Analyst Intern (Jan–Mar 2026)
   M&A boutique. Built financial models for forecasting, scenario and variance analysis.
   Industry and macro analysis for valuation assumptions. Prepared client-ready presentations.
   Helped lead M&A transactions: financial analysis, teasers, CIMs, pitchbooks.

3. Lido Advisors, Los Angeles/San Diego — Wealth Management Intern (Jun–Aug 2025)
   $30B AUM wealth management firm. Built ML-driven portfolio reweighting models with ESG tilts.
   Python-based ESG framework for S&P 500 sector allocation. Client research and prospecting.
   Presented portfolio recommendations to high-net-worth clients.

4. Visionaries 777, Hong Kong — Financial Analyst Intern (Aug–Dec 2023)
   AR/VR tech startup (clients: Lego, Cartier). Built 3-year financial forecast and valuation model.
   Automated financial models for investor reporting. CEO pitch decks for stakeholder presentations.

5. BNP Paribas, Nice — Wealth Management Intern (May–Jul 2022)
   Weekly market reports identifying trends and investment opportunities.
   DCF-based valuations and financial analysis for SMEs to support lending/advisory decisions.

TECHNICAL SKILLS:
Python, SQL, MATLAB, R, Tableau, SAS, Excel (advanced), PowerPoint

PROGRAMMING PROJECTS:
- SQL/Python pipeline with LA Capital Management: analysed analyst targets, corporate events,
  and prices to design and backtest market-neutral event-driven trading strategies
- Quantitative Asset Management: designed and backtested long-short equity strategy exploiting
  mispricing in R&D vs advertising intensity

PERSONAL:
- Competitive swimmer at national level 2011–2023 (French National Championships, FFN Golden Tour,
  CCS Nationals). Instilled discipline, performance under pressure, long-term commitment.
- Interests: marathons, scuba diving, rock-climbing, surfing, literature, travel
"""


# ─────────────────────────────────────────────────────────────────
# COVER LETTER EXAMPLES — Claude learns Emilie's voice from these
# ─────────────────────────────────────────────────────────────────

COVER_LETTER_EXAMPLES = """
EXAMPLE 1 — SocGen VIE M&A (strong example of her voice):
---
Dear Hiring Manager,
I am writing to express my strong interest in the VIE M&A Analyst position within Société Générale's
Client Advisory Team. As a Financial Engineering student at UCLA Anderson with cross-border experience
spanning France, Hong Kong, and the United States, I am eager to contribute to SG's M&A advisory
capabilities while developing expertise in valuation and deal execution.
My quantitative training directly aligns with your analytical requirements. At UCLA, I've mastered
financial modeling, derivative pricing, and risk management—skills essential for building DCF models,
LBO analyses, and pro forma projections. During my summer internship at Lido Advisors, I developed
portfolio optimization models and valuation frameworks, strengthening my ability to construct complex
financial analyses under tight deadlines.
At Visionaries 777, I created dynamic financial models and CEO pitch decks for stakeholder
presentations, experience that translates directly to preparing client materials for live transactions
and pitch meetings. My work at BNP Paribas further reinforced my understanding of financial analysis
and market dynamics within the French banking ecosystem.
What distinguishes my candidacy is my international perspective. Native fluency in French and English,
combined with fluent Chinese, positions me to excel in cross-border M&A situations and coordinate
seamlessly between SG's New York office and global teams in Paris, London, and Asia.
Thank you for your consideration.
Sincerely, Emilie Quillet
---

EXAMPLE 2 — SocGen VIE IB (shows how she handles IB-specific tone):
---
Dear Hiring Manager,
I'm drawn to this V.I.E. Investment Banking position because it sits at the intersection of strategic
corporate finance and M&A execution—exactly where I want to build my career. The Advisory & Financing
Group's mandate is unique: you're not just executing transactions, you're serving as strategic CFO
advisor to corporations across their most critical decisions. What excites me most is the analytical
depth—building bespoke frameworks for capital allocation, valuation drivers, and capital structure
optimisation rather than cookie-cutter pitches.
My financial modelling foundation directly supports this mandate. At UCLA Anderson, I've mastered DCF
valuation, capital structure analysis, and derivative pricing. At Lido Advisors, I built portfolio
optimisation models and valuation frameworks under tight deadlines. At Visionaries 777, I created
automated financial models and CEO pitch decks that translated complex scenarios for stakeholders.
The cross-product coordination aspect is particularly appealing; working at the nexus of M&A, ECM,
credit ratings, and derivatives means I'd be constantly learning how different financing solutions
interconnect.
What distinguishes my candidacy is my cross-border perspective. Native French and English fluency,
combined with fluent Chinese, positions me to contribute immediately to international mandates.
Thank you for your consideration.
Sincerely, Emilie Quillet
---

EXAMPLE 3 — Keensight Capital PE (shows how she adapts for PE):
---
Dear Hiring Manager,
Keensight Capital's focus on fast-growing Healthcare and Technology companies resonates with my
quantitative training and passion for analysing high-potential investments. As a Financial Engineering
student at UCLA Anderson with cross-border experience spanning France, Hong Kong, and the United
States, I am eager to support your investment team's evaluation and execution activities.
My UCLA coursework in financial modelling, derivative markets, and data analysis provides the technical
foundation essential for building operating models and assessing company performance. At Lido Advisors,
I developed portfolio optimisation models and valuation frameworks. At Visionaries 777, I created
dynamic financial models for AR/VR solutions and presented CEO pitch decks to stakeholders.
What distinguishes my candidacy is my cross-border perspective. Having lived and worked across France,
Hong Kong, and the US, I understand how to navigate different business cultures. Native French and
English fluency positions me to coordinate seamlessly between offices.
Thank you for your consideration.
Sincerely, Emilie Quillet
---

EXAMPLE 4 — Crédit Agricole Coverage (shows relationship banking tone):
---
Dear Hiring Manager,
I'm drawn to this Coverage Analyst position because it offers exactly what I am seeking: hands-on
exposure to relationship banking and deal execution within a global French institution. The role's
emphasis on multi-product coordination and strategic client mapping aligns perfectly with my analytical
training and relationship-building experience from BNP Paribas and Visionaries 777.
At BNP Paribas, I assisted senior advisors by preparing weekly market reports and analysing client
financials. At Visionaries 777, I developed financial models and CEO pitch decks for stakeholder
presentations, strengthening my ability to translate complex analyses into compelling client materials.
My UCLA training in financial modelling and derivatives provides the foundation for analysing
risk-return metrics and supporting credit requests.
What distinguishes my candidacy is my cross-border perspective. Native French and English fluency,
combined with fluent Chinese, positions me to coordinate seamlessly between CACIB's Americas platform
and global offices while fostering international client relationships.
Thank you for your consideration.
Sincerely, Emilie Quillet
---
"""


# ─────────────────────────────────────────────────────────────────
# TONE RULES BY FIRM TYPE
# ─────────────────────────────────────────────────────────────────

def get_tone_for_company(company: str, title: str) -> str:
    c = company.lower()
    t = title.lower()

    if any(x in c for x in ["rothschild", "lazard", "perella", "evercore", "centerview"]):
        return (
            "TONE: Elite boutique. Formal, precise, understated. Let credentials speak — "
            "no enthusiasm-signalling, no exclamation marks. Every word must earn its place. "
            "Lead with intellectual substance, not eagerness."
        )
    elif any(x in c for x in ["eurazeo", "tikehau", "ardian", "kkr", "carlyle", "blackstone", "apollo"]):
        return (
            "TONE: Private equity. Sharp and investor-minded. Lead with quantitative instincts. "
            "Show you think about returns, not just process. Reference modelling and analytical rigour prominently."
        )
    elif any(x in c for x in ["deloitte", "kpmg", "pwc", "ey", "bearingpoint", "mckinsey", "bcg", "bain"]):
        return (
            "TONE: Consulting / Big 4. Professional but warm. Emphasise team collaboration, "
            "client communication, and cross-functional work. Analytical rigour paired with "
            "clear communication is the key message."
        )
    elif any(x in c for x in ["bnp", "societe", "société", "credit agricole", "crédit agricole",
                                "natixis", "cacib", "sg ", "soc gen"]):
        return (
            "TONE: French bank. Collegial and professional. Emphasise the French-language advantage "
            "and understanding of French banking culture. Cross-border coordination is a strong card."
        )
    elif any(x in c for x in ["jp morgan", "goldman", "morgan stanley", "bank of america", "citi", "barclays"]):
        return (
            "TONE: Bulge bracket. Confident and high-energy. Show intellectual curiosity and "
            "drive. These firms want ambition paired with analytical foundation."
        )
    else:
        return (
            "TONE: Direct and confident. Match the energy of the JD. "
            "If the JD is formal, be formal. If conversational, be slightly warmer. "
            "Always lead with something specific to this company and role."
        )


# ─────────────────────────────────────────────────────────────────
# SETUP
# ─────────────────────────────────────────────────────────────────

def setup_directories():
    for d in [OUTPUT_DIR, COVER_LETTERS_DIR, MODIFIED_CVS_DIR]:
        d.mkdir(parents=True, exist_ok=True)
    if not CSV_LOG_FILE.exists():
        with open(CSV_LOG_FILE, "w", newline="", encoding="utf-8") as f:
            csv.writer(f).writerow(CSV_HEADERS)
    if not SEEN_JOBS_FILE.exists():
        SEEN_JOBS_FILE.write_text("[]")


def load_seen_jobs() -> set:
    return set(json.loads(SEEN_JOBS_FILE.read_text()))


def save_seen_jobs(seen: set):
    SEEN_JOBS_FILE.write_text(json.dumps(list(seen)))


# ─────────────────────────────────────────────────────────────────
# JOB SCRAPING
# ─────────────────────────────────────────────────────────────────

def scrape_all_jobs() -> list[dict]:
    all_jobs = []
    seen_urls = set()

    for term in SEARCH_CONFIG["search_terms"]:
        print(f"  Searching: {term}...")
        try:
            jobs = scrape_jobs(
                site_name=SEARCH_CONFIG["sites"],
                search_term=term,
                location=SEARCH_CONFIG["location"],
                results_wanted=SEARCH_CONFIG["results_per_search"],
                hours_old=SEARCH_CONFIG["hours_old"],
                country_indeed="Luxembourg",
            )
            for _, row in jobs.iterrows():
                url = str(row.get("job_url", "")).strip()
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    all_jobs.append({
                        "title": str(row.get("title", "Unknown")).strip(),
                        "company": str(row.get("company", "Unknown")).strip(),
                        "location": str(row.get("location", "")).strip(),
                        "site": str(row.get("site", "")).strip(),
                        "url": url,
                        "date_posted": str(row.get("date_posted", "")).strip(),
                        "description": str(row.get("description", ""))[:4000],
                    })
            time.sleep(2)
        except Exception as e:
            print(f"  Warning: '{term}' search failed — {e}")

    print(f"  Total unique jobs found: {len(all_jobs)}")
    return all_jobs


# ─────────────────────────────────────────────────────────────────
# COVER LETTER GENERATION
# ─────────────────────────────────────────────────────────────────

def generate_cover_letter(job: dict, client: anthropic.Anthropic) -> str:
    tone = get_tone_for_company(job["company"], job["title"])

    prompt = f"""You are writing a cover letter for Emilie Quillet. Study her profile and voice carefully.

═══════════════════════
EMILIE'S PROFILE
═══════════════════════
{EMILIE_PROFILE}

═══════════════════════
HER VOICE — learn from these real examples she wrote and approved:
═══════════════════════
{COVER_LETTER_EXAMPLES}

═══════════════════════
THE JOB
═══════════════════════
Job Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
Job Description:
{job['description']}

═══════════════════════
TONE INSTRUCTION FOR THIS FIRM
═══════════════════════
{tone}

═══════════════════════
STRICT RULES — follow all of these
═══════════════════════
STRUCTURE (3 short paragraphs max):
1. Opening: What specifically draws her to THIS role at THIS company. Reference something real
   from the JD — a mandate, a team structure, a specific product. Never generic.
2. Middle: Connect 2 of her most relevant experiences directly to the role's key requirements.
   Be specific — name the company, the output, the skill. No vague claims.
3. Close: What she uniquely brings (multilingual, cross-border, quant + interpersonal).
   Confident sign-off — not "I hope to hear from you."

NEVER write:
- "I am writing to express my interest" or any variation
- "I am passionate about finance" — too generic
- "I would be a great fit" — let the evidence speak
- "Please find attached my CV"
- Any sentence that could appear unchanged in someone else's letter
- Bullet points inside the letter
- More than 280 words total

ALWAYS:
- Address to "Dear Hiring Manager,"
- Sign off: "Sincerely, Emilie Quillet"
- Write in British English (she studied in France/UK context, based in Luxembourg)
- The swimming background can appear sparingly only if it genuinely fits the role's language
  (discipline, performance under pressure) — never forced

Return only the cover letter text. No preamble, no explanation."""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=1000,
        messages=[{"role": "user", "content": prompt}]
    )
    return message.content[0].text.strip()


def save_cover_letter(job: dict, letter: str) -> str:
    filename = make_filename(job, suffix="CoverLetter", ext="txt")
    (COVER_LETTERS_DIR / filename).write_text(letter, encoding="utf-8")
    return filename


# ─────────────────────────────────────────────────────────────────
# CV KEYWORD INJECTION
# ─────────────────────────────────────────────────────────────────

def extract_ats_keywords(job: dict, client: anthropic.Anthropic) -> list[str]:
    """Ask Claude to extract the key ATS keywords from the JD."""
    prompt = f"""Extract the most important ATS (applicant tracking system) keywords from this job description.
These are the specific skills, tools, qualifications, and role-specific terms that an ATS would scan for.

Job Title: {job['title']}
Company: {job['company']}
Job Description:
{job['description']}

Return ONLY a comma-separated list of keywords and short phrases. No explanation, no numbering.
Focus on: technical skills, software tools, financial methods, role-specific terminology, qualifications.
Maximum 40 keywords. Example format: DCF, LBO, financial modelling, Excel, PowerPoint, due diligence, M&A"""

    message = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=300,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = message.content[0].text.strip()
    keywords = [kw.strip() for kw in raw.split(",") if kw.strip()]
    return keywords


def inject_keywords_into_cv(job: dict, keywords: list[str]) -> str:
    """
    Opens the base CV .docx, appends a hidden paragraph at the bottom with
    ATS keywords in white font size 1 (invisible to human eye, readable by ATS).
    Saves as a new file per job application.
    """
    if not BASE_CV_PATH.exists():
        raise FileNotFoundError(
            f"Base CV not found at {BASE_CV_PATH}. "
            "Update BASE_CV_PATH at the top of this script to point to your CV_Emilie_Quillet_v15.docx"
        )

    doc = Document(BASE_CV_PATH)

    # Build the invisible keyword string
    keyword_text = " | ".join(keywords)

    # Add a new paragraph at the end of the document
    hidden_para = doc.add_paragraph()
    hidden_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = hidden_para.add_run(keyword_text)

    # Make it invisible: white font, size 1pt
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # pure white
    run.font.size = Pt(1)

    # Save modified CV
    filename = make_filename(job, suffix="CV", ext="docx")
    filepath = MODIFIED_CVS_DIR / filename
    doc.save(filepath)

    return filename


# ─────────────────────────────────────────────────────────────────
# UTILITIES
# ─────────────────────────────────────────────────────────────────

def make_filename(job: dict, suffix: str, ext: str) -> str:
    date = datetime.now().strftime("%Y%m%d")
    safe_company = re.sub(r'[^\w\s-]', '', job["company"])[:25].strip().replace(" ", "_")
    safe_title = re.sub(r'[^\w\s-]', '', job["title"])[:25].strip().replace(" ", "_")
    return f"{date}_{safe_company}_{safe_title}_{suffix}.{ext}"


def log_to_csv(job: dict, cover_letter_file: str, cv_file: str):
    with open(CSV_LOG_FILE, "a", newline="", encoding="utf-8") as f:
        csv.writer(f).writerow([
            datetime.now().strftime("%Y-%m-%d %H:%M"),
            job["title"],
            job["company"],
            job["location"],
            job["site"],
            job["url"],
            job["date_posted"],
            str(COVER_LETTERS_DIR / cover_letter_file),
            str(MODIFIED_CVS_DIR / cv_file),
            "No",   # Applied? — you update this
            "",     # Response
            ""      # Notes
        ])


# ─────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────

def run():
    print("\n" + "="*50)
    print("JOB AUTOMATOR — Emilie Quillet")
    print(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    print("="*50 + "\n")

    if not ANTHROPIC_API_KEY:
        print("ERROR: ANTHROPIC_API_KEY environment variable not set.")
        print("Run: export ANTHROPIC_API_KEY='sk-ant-your-key-here'")
        return

    if not BASE_CV_PATH.exists():
        print(f"ERROR: CV not found at {BASE_CV_PATH}")
        print("Update BASE_CV_PATH at the top of this script.")
        return

    setup_directories()
    seen_jobs = load_seen_jobs()
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    print("Step 1: Scraping jobs...")
    all_jobs = scrape_all_jobs()

    new_jobs = [j for j in all_jobs if j["url"] not in seen_jobs]
    print(f"\nStep 2: {len(new_jobs)} new jobs to process\n")

    if not new_jobs:
        print("Nothing new today. Check again tomorrow.")
        return

    success_count = 0
    for i, job in enumerate(new_jobs, 1):
        print(f"[{i}/{len(new_jobs)}] {job['title']} — {job['company']}")

        try:
            # Generate cover letter
            letter = generate_cover_letter(job, client)
            cl_file = save_cover_letter(job, letter)
            print(f"  ✓ Cover letter: {cl_file}")

            # Extract keywords and inject into CV
            keywords = extract_ats_keywords(job, client)
            cv_file = inject_keywords_into_cv(job, keywords)
            print(f"  ✓ Modified CV:  {cv_file}")
            print(f"  ✓ Keywords injected ({len(keywords)}): {', '.join(keywords[:8])}...")

            # Log to CSV
            log_to_csv(job, cl_file, cv_file)
            seen_jobs.add(job["url"])
            success_count += 1

            time.sleep(1.5)  # pace API calls

        except Exception as e:
            print(f"  ✗ Failed: {e}")

    save_seen_jobs(seen_jobs)

    print(f"\n{'='*50}")
    print(f"Done. {success_count}/{len(new_jobs)} jobs processed successfully.")
    print(f"\nYour files are in: {OUTPUT_DIR}")
    print(f"Open job_log.csv to review and apply.")
    print("="*50 + "\n")


if __name__ == "__main__":
    run()


# ═══════════════════════════════════════════════════════════════
# SETUP GUIDE
# ═══════════════════════════════════════════════════════════════
#
# ONE-TIME SETUP (do this once on your personal PC):
#
# 1. INSTALL PYTHON
#    python.org → download version 3.11 or higher
#    Verify: open Terminal → type: python --version
#
# 2. INSTALL LIBRARIES
#    Open Terminal and run:
#      pip install python-jobspy
#      pip install anthropic
#      pip install python-docx
#
# 3. SAVE YOUR CV
#    Save CV_Emilie_Quillet_v15.docx in the SAME folder as this script
#    OR update the BASE_CV_PATH line at the top to point to wherever it is
#    Example: BASE_CV_PATH = Path("/Users/emilie/Documents/CV_Emilie_Quillet_v15.docx")
#
# 4. GET YOUR ANTHROPIC API KEY
#    Go to console.anthropic.com → create account → API Keys → create key
#    Copy the key (starts with sk-ant-...)
#    Cost: ~$0.01 per job (cover letter + keyword extraction). 20 jobs/day = ~$6/month
#
# 5. SET THE API KEY (Mac/Linux)
#    Open Terminal and run:
#      export ANTHROPIC_API_KEY="sk-ant-your-key-here"
#    To make permanent, add that line to ~/.zshrc or ~/.bashrc
#
#    On Windows:
#      set ANTHROPIC_API_KEY=sk-ant-your-key-here
#
# 6. RUN IT
#    In Terminal, navigate to the script folder:
#      cd ~/Desktop/job_search
#    Then:
#      python job_automator_v2.py
#
# 7. SCHEDULE IT (runs every day at 7am — Mac/Linux)
#    In Terminal:
#      crontab -e
#    Add this line (update the path):
#      0 7 * * * python3 /Users/emilie/Desktop/job_search/job_automator_v2.py
#
# YOUR DAILY ROUTINE (10-15 min):
#   1. Open ~/Desktop/job_applications/job_log.csv
#   2. Review new rows — columns tell you title, company, URL, file paths
#   3. Click the cover letter path → read and tweak if needed
#   4. Click the modified CV path → that's the one to upload (has hidden keywords)
#   5. Apply via the URL in column F
#   6. Mark Applied = Yes in the sheet
#
# TROUBLESHOOTING:
#   "ModuleNotFoundError"   → re-run: pip install python-jobspy anthropic python-docx
#   "AuthenticationError"   → ANTHROPIC_API_KEY not set correctly, recheck step 5
#   "FileNotFoundError CV"  → update BASE_CV_PATH at top of script to correct path
#   JobSpy returns 0 results → LinkedIn blocking temporarily, try again in a few hours
#                              or reduce results_per_search to 10
# ═══════════════════════════════════════════════════════════════
